import json
import os
import pathlib
import re
import zipfile
import xml.etree.ElementTree
import pandas as pd
import requests as requests
from docx import Document

"""Константы"""
USERS = ('Бутакова', 'Винтова', 'Грибанова', 'Чернова')
WORD_NAMESPACE = '{http://schemas.openxmlformats.org/wordprocessingml/2006/main}'
PARA = WORD_NAMESPACE + 'p'
TEXT = WORD_NAMESPACE + 't'
TABLE = WORD_NAMESPACE + 'tbl'
ROW = WORD_NAMESPACE + 'tr'
CELL = WORD_NAMESPACE + 'tc'

TEXT_FILTERED = []


class ParserFile:
    @staticmethod
    def open_file(file):
        try:
            with zipfile.ZipFile(file) as docx:
                return xml.etree.ElementTree.XML(docx.read('word/document.xml'))
        except zipfile.BadZipFile:
            return False

    @staticmethod
    def get_all_text(file):
        list_text = []
        for text in file.iter(PARA):
            item = ''.join(node.text for node in text.iter(TEXT))
            if item == '' or item == ' ' or item.__len__() < 2:
                continue
            else:
                list_text.append(item)
        return list_text

    @staticmethod
    def find_tasks(text):
        support_param = True
        for el in text:
            if support_param:
                if el[:2] == '2.':
                    index_start_task = text.index(el)
                    support_param = False
            if el[:25] == 'Настоящий Отчет по резуль':
                index_stop_task = text.index(el)
                break
        return text[index_start_task:index_stop_task]

    @staticmethod
    def clear_num_page_in_task(task):
        task = task[::-1]
        index = 0
        for el in task:
            try:
                int(el)
                index += 1
            except ValueError:
                break
        return task[index:][::-1]

    @staticmethod
    def tree_tasks(tasks):
        LIST_TASK_FILTERED = []
        task_tree_list = {}
        for task in tasks:
            task_split = task.split('.')
            index = 0

            for el in task_split:
                try:
                    int(el)
                    index += 1
                except ValueError:
                    break

            if index == 1:
                task_tree_list[task_split[0]] = task[2:]
                LIST_TASK_FILTERED.append(task[2:])
            if index == 2:
                task_tree_list[f"{task_split[0]}{task_split[1]}"] = task[4:]
                LIST_TASK_FILTERED.append(task[4:])
            if index == 3:
                task_tree_list[f"{task_split[0]}{task_split[1]}{task_split[2]}"] = task[6:]
                LIST_TASK_FILTERED.append(task[6:])
            if index == 4:
                task_tree_list[f"{task_split[0]}{task_split[1]}{task_split[2]}{task_split[3]}"] = task[8:]
                LIST_TASK_FILTERED.append(task[8:])
            if index == 5:
                task_tree_list[f"{task_split[0]}{task_split[1]}{task_split[2]}{task_split[3]}{task_split[4]}"] = task[
                                                                                                                 10:]
                LIST_TASK_FILTERED.append(task[10:])

        return task_tree_list, LIST_TASK_FILTERED

    @staticmethod
    def all_text_in_doc(file):
        document = Document(file)
        TEXt = []
        for parag in document.paragraphs:
            if parag.text == '' or parag.text == '\n':
                continue
            else:
                TEXt.append(parag.text)
        for text in TEXt:
            if text[:76] == 'иные нормативно-правовые акты Российской Федерации, необходимые для анализа.':
                index = TEXt.index(text)
                return TEXt[index:]

    @staticmethod
    def task_text(start_task, stop_task, all_text):
        if stop_task:
            if start_task[0] == '.':
                start_task = start_task[1:]
            if stop_task[0] == '.':
                stop_task = stop_task[1:]
            for text in all_text:
                text_clear = ParserFile.clear_num_page_in_task(text)
                if text_clear.strip().replace("\xa0", " ").replace("\n", "").replace(' ', '').replace(' ', '').replace(
                        '…', '').replace('.', '') == start_task.strip().replace("\xa0", " ").replace(' ', '').replace(
                        ' ', '').replace('…', '').replace('.', '') or text_clear.strip().replace("\xa0", " ").replace(
                        "\n", "").replace(' ', '').replace(' ', '').replace('…', '').replace('.',
                                                                                             '') == start_task.strip().replace(
                        "\xa0", "").replace(' ', '').replace(' ', '').replace('…', '').replace('.', ''):
                    index_start = all_text.index(text)
                elif text_clear.strip().replace("\xa0", " ").replace("\n", "").replace(' ', '').replace(' ',
                                                                                                        '').replace('…',
                                                                                                                    '').replace(
                        '.', '') == stop_task.strip().replace("\xa0", " ").replace(' ', '').replace(' ', '').replace(
                        '…', '').replace('.', '') or text_clear.strip().replace("\xa0", " ").replace("\n", "").replace(
                        ' ', '').replace(' ', '').replace('…', '').replace('.', '') == start_task.strip().replace(
                        "\xa0", "").replace(' ', '').replace(' ', '').replace('…', '').replace('.', ''):
                    index_stop = all_text.index(text)
                    break
                else:
                    continue
            return all_text[index_start + 1:index_stop], index_stop
        else:
            if start_task[0] == '.':
                start_task = start_task[1:]
            for text in all_text:
                if text.strip().replace("\xa0", " ") == start_task.strip().replace("\xa0", " "):
                    index_start = all_text.index(text)
                else:
                    continue
            return all_text[index_start + 1:], index_start

    @staticmethod
    def get_tables(file):
        document = Document(file)
        for table in document.tables:
            print(1)
            for row in table.rows:
                print(row.text)
                for col in table.columns:
                    print(col.text)
        print(1)


def work_in_file(dir, user):
    try:
        for currentFile in dir.iterdir():
            treeFile = ParserFile.open_file(currentFile)  # Получаем дерево обьектов файла
            if not treeFile:
                continue
            all_text = ParserFile.get_all_text(treeFile)  # Получаем весь текст файла
            all_tasks_notclear = ParserFile.find_tasks(all_text)  # Поиск всех задач из файла
            tasks = []
            for task in all_tasks_notclear:
                tasks.append(ParserFile.clear_num_page_in_task(task))
            task_tree = ParserFile.tree_tasks(tasks)  # Создание дерева задач

            # response = requests.post('http://94.26.245.131/api-create-task/',
            #                          data={
            #                              "user": user,
            #                              "task": json.dumps(task_tree[0]),
            #                              'project': currentFile.stem.replace('Отчет', "")
            #                          })
            text_filtered = ParserFile.all_text_in_doc(currentFile)

            index = 0
            i = 0
            while i < task_tree[1].__len__() - 1:
                try:
                    task_text = ParserFile.task_text(task_tree[1][i], task_tree[1][i + 1], text_filtered[index:])
                    index += task_text[1]
                    i += 1
                except IndexError:
                    task_text = ParserFile.task_text(task_tree[1][i], False, text_filtered[index:])
                    i += 1

                glue_text = ''
                for text in task_text[0]:
                    glue_text += f'\n\t{text}'

                # response = requests.post('http://94.26.245.131/api-create-text/',
                #                          data={
                #                              "user": user,
                #                              "task": task_tree[1][i],
                #                              "text": glue_text,
                #                              'project': currentFile.stem.replace('Отчет', "")
                #                          })
                print(1)
    except KeyError:
        pass


def start():
    for user in USERS:
        currentDirectory = pathlib.Path(f'./files/{user}/')  # Открываем директорию с файлами
        work_in_file(currentDirectory, user)  # Запускаем цикл по каждому файлу в директории


if __name__ == '__main__':
    start()
